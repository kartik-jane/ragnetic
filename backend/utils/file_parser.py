import os
from pypdf import PdfReader
import docx

def parse_file(file_path):
    """
    Detect file type & extract text accordingly.
    Supports: PDF, DOCX, TXT
    """

    ext = os.path.splitext(file_path)[-1].lower()
    print(f"[PARSE] Parsing file with extension: {ext}, path: {file_path}")

    if ext == ".pdf":
        return parse_pdf(file_path)

    elif ext == ".docx":
        return parse_docx(file_path)

    elif ext == ".txt":
        return parse_txt(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def parse_pdf(path):
    print(f"[PARSE] Parsing PDF: {path}")
    try:
        text = ""
        reader = PdfReader(path)
        print(f"[PARSE] PDF has {len(reader.pages)} pages")
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            print(f"[PARSE] Page {i+1} extracted {len(page_text) if page_text else 0} characters")
            text += page_text + "\n" if page_text else ""
        print(f"[PARSE] PDF total text: {len(text)} characters")
        return text
    except Exception as e:
        print(f"[PARSE] PDF parsing error: {str(e)}")
        raise


def parse_docx(path):
    print(f"[PARSE] Parsing DOCX: {path}")
    try:
        doc = docx.Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        print(f"[PARSE] DOCX extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
        return text
    except Exception as e:
        print(f"[PARSE] DOCX parsing error: {str(e)}")
        raise


def parse_txt(path):
    print(f"[PARSE] Parsing TXT: {path}")
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        print(f"[PARSE] TXT extracted {len(text)} characters")
        return text
    except Exception as e:
        print(f"[PARSE] TXT parsing error: {str(e)}")
        raise
